import os
import re
from pypdf import PdfReader
from docx import Document

class SimpleDocument:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata if metadata else {}

def load_document(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    metadata = {"source": os.path.basename(file_path)}
    
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return [SimpleDocument(content, metadata)]
    
    elif ext == '.pdf':
        reader = PdfReader(file_path)
        documents = []
        for page_num, page in enumerate(reader.pages):
            content = page.extract_text()
            if content:
                page_metadata = metadata.copy()
                page_metadata["page"] = page_num + 1
                documents.append(SimpleDocument(content, page_metadata))
        return documents
    
    elif ext == '.docx':
        doc = Document(file_path)
        content = "\n".join([para.text for para in doc.paragraphs])
        return [SimpleDocument(content, metadata)]
    
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def load_documents_from_folder(folder_path):
    documents = []
    supported_extensions = ('.txt', '.pdf', '.docx')
    
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(supported_extensions):
            file_path = os.path.join(folder_path, filename)
            try:
                docs = load_document(file_path)
                documents.extend(docs)
                print(f"Loaded {filename}")
            except Exception as e:
                print(f"Failed to load {filename}: {e}")
    
    return documents

def split_documents(documents, chunk_size=1000, chunk_overlap=200):
    chunks = []
    
    for doc in documents:
        text = doc.page_content
        metadata = doc.metadata
        
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + chunk_size, text_length)
            
            if end < text_length:
                last_period = text.rfind('。', start, end)
                last_newline = text.rfind('\n', start, end)
                if last_period > start + chunk_overlap:
                    end = last_period + 1
                elif last_newline > start + chunk_overlap:
                    end = last_newline + 1
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunk_metadata = metadata.copy()
                chunk_metadata["chunk_start"] = start
                chunks.append(SimpleDocument(chunk_text, chunk_metadata))
            
            if start + chunk_size >= text_length:
                break
            
            start = end - chunk_overlap
    
    return chunks